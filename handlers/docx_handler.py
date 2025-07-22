"""
DOCX Metadata Scrubber for rMeta

Uses python-docx to copy paragraph contents into a new document,
which excludes original metadata and revision history.
"""

import docx
import os

supported_extensions = {"docx"}


def scrub(file_path):
    """
    Scrubs metadata from a DOCX file by copying contents to a clean document.

    Args:
        file_path (str): Path to the input DOCX file.
    """
    try:
        doc = docx.Document(file_path)
        new_doc = docx.Document()

        # Copy plain paragraph text (no formatting, no metadata)
        for para in doc.paragraphs:
            new_doc.add_paragraph(para.text)

        cleaned_path = file_path.replace(".docx", "_cleaned.docx")
        new_doc.save(cleaned_path)
        os.replace(cleaned_path, file_path)
    except Exception as e:
        print(f"Error scrubbing DOCX metadata: {e}")
